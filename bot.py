import os
import telebot
from openai import OpenAI
import docx
from datetime import datetime
import json

# טעינת משתני הסביבה
TOKEN = os.environ.get('TELEGRAM_TOKEN')
OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY')

bot = telebot.TeleBot(TOKEN)
client = OpenAI(api_key=OPENAI_API_KEY)

# מילון לשמירת מצב השיחה של המשתמשים
user_states = {}

def get_missing_fields_prompt(text):
    """שולח את הטקסט ל-OpenAI כדי לבדוק אם חסרים פרטים"""
    prompt = f"""
    אתה עוזר מקצועי של חברת "סטאר מהנדסים" שתפקידו לעבד נתונים מהשטח ולהכין דוח פיקוח עליון מובנה.
    נתח את הטקסט הבא ששלח המהנדס מהשטח וזהה איזה מהפרטים הבאים חסרים:
    1. project_num (מספר פרויקט)
    2. letter_num (מספר מכתב)
    3. client_name (לכבוד - שם הלקוח)
    4. contact_person (לידי - איש קשר)
    5. email (במייל - כתובת אימייל)
    6. structure_name (שם המבנה / הפרויקט)
    7. supervisor_name (שם המפקח באתר)
    8. contractor_name (נציגי הביצוע / קבלן)
    9. author_initials (ראשי תיבות של כותב הדוח)
    10. description (תיאור מצב העבודה הנוכחי)

    שים לב לחוקים הבאים:
    - נוכח מטעם סטאר מהנדסים הוא תמיד "הח"מ" כברירת מחדל, אלא אם צוין אחרת במפורש, לכן אל תסמן אותו כחסר.
    - אל תמציא פרטים! אם פרט לא מופיע, הוא חסר.

    החזר אך ורק תשובת JSON תקינה בפורמט הבא, ללא שום טקסט נוסף לפני או אחרי:
    {{
        "missing_fields": ["שם השדה החסר בעברית", "עוד שדה חסר בעברית"],
        "extracted_data": {{
            "project_num": "הערך שנמצא או null",
            "letter_num": "הערך שנמצא או null",
            "client_name": "הערך שנמצא או null",
            "contact_person": "הערך שנמצא או null",
            "email": "הערך שנמצא או null",
            "structure_name": "הערך שנמצא או null",
            "supervisor_name": "הערך שנמצא או null",
            "contractor_name": "הערך שנמצא או null",
            "author_initials": "הערך שנמצא או null",
            "description": "הערך שנמצא או null"
        }}
    }}

    הטקסט מהשטח:
    "{text}"
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            temperature=0
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        print(f"Error calling OpenAI: {e}")
        return None

def improve_description(data):
    """משתמש ב-OpenAI כדי לנסח את תיאור המצב בשפה הנדסית גבוהה"""
    prompt = f"""
    אתה מהנדס מבנים בכיר ב"סטאר מהנדסים". נסח מחדש את תיאור מצב העבודה הבא שנכתב כהערות קצרות מהשטח, 
    והפוך אותו לפסקה מנוסחת היטב בשפה הנדסית רשמית, מקצועית ומדויקת המתאימה לדוח פיקוח עליון.
    
    הנתונים הגולמיים מהשטח: "{data.get('description')}"
    """
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3
        )
        return response.choices[0].message.content.strip()
    except:
        return data.get('description')

def create_report(data, template_path="template.docx", output_path="output.docx"):
    """משבץ את הנתונים לתוך תבנית הוורד הקבועה"""
    doc = docx.Document(template_path)
    
    # שיפור הניסוח של התיאור
    professional_description = improve_description(data)
    
    replacements = {
        "[מספר פרויקט]": str(data.get("project_num", "") or ""),
        "[מספר מכתב]": str(data.get("letter_num", "") or ""),
        "[התאריך של היום]": datetime.now().strftime("%d/%m/%Y"),
        "[שם הלקוח]": str(data.get("client_name", "") or ""),
        "[איש קשר]": str(data.get("contact_person", "") or ""),
        "[כתובת אימייל]": str(data.get("email", "") or ""),
        "[שם המבנה / הפרויקט]": str(data.get("structure_name", "") or ""),
        "[שם המפקח באתר]": str(data.get("supervisor_name", "") or ""),
        "[נציגי הביצוע]": str(data.get("contractor_name", "") or ""),
        "[ראשי תיבות של כותב הדוח]": str(data.get("author_initials", "") or ""),
        "[תיאור מצב העבודה הנוכחי]": professional_description,
        "[נוכח מטעם סטאר]": "הח\"מ"
    }
    
    # החלפה בפסקאות רגילות
    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
                
    # החלפה בתוך טבלאות (אם יש בתבנית)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in replacements.items():
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)
                            
    doc.save(output_path)
    return output_path

@bot.message_with_type_handler(commands=['start', 'help'])
def send_welcome(message):
    bot.reply_to(message, "שלום אביב! שלח לי את הנתונים מהסיור בשטח (בטקסט או בהקלטה קולית), ואני אכין לך את דוח הפיקוח העליון בוורד.")

@bot.message_handler(func=lambda message: True)
def handle_text(message):
    chat_id = message.chat.id
    text = message.text

    # אם המשתמש נמצא בתהליך השלמת פרטים חסרים
    if chat_id in user_states and user_states[chat_id].get("waiting_for"):
        state = user_states[chat_id]
        missing_field_key = state["waiting_for"]
        
        # שמירת המידע שהמשתמש הזין
        state["data"][missing_field_key] = text
        state["missing_list"].pop(0) # הסרת השדה שהושלם מהרשימה
        
        # בדיקה אם יש עוד שדות חסרים
        if state["missing_list"]:
            next_field = state["missing_list"][0]
            # מיפוי קל כדי לשאול בצורה יפה
            prompts_map = {
                "project_num": "מספר פרויקט", "letter_num": "מספר מכתב",
                "client_name": "שם הלקוח (לכבוד)", "contact_person": "איש קשר (לידי)",
                "email": "כתובת אימייל (במייל)", "structure_name": "שם המבנה / הפרויקט",
                "supervisor_name": "שם המפקח באתר", "contractor_name": "נציגי הביצוע / קבלן",
                "author_initials": "ראשי תיבות של כותב הדוח"
            }
            field_display = prompts_map.get(next_field, next_field)
            state["waiting_for"] = next_field
            bot.send_message(chat_id, f"מעולה. עכשיו, מהו **{field_display}**?")
            return
        else:
            # הכל הושלם! מייצרים דוח
            bot.send_message(chat_id, "כל הפרטים נאספו בהצלחה! מייצר עבורך את קובץ הוורד...")
            try:
                doc_file = create_report(state["data"])
                with open(doc_file, 'rb') as f:
                    bot.send_document(chat_id, f, caption="הנה דוח הפיקוח העליון המוכן שלך! 📝")
                user_states.pop(chat_id, None) # איפוס מצב
            except Exception as e:
                bot.send_message(chat_id, f"אירעה שגיאה ביצירת הקובץ: {e}")
            return

    # תהליך התחלתי - ניתוח הטקסט הראשוני
    bot.send_message(chat_id, "מנתח את הנתונים ששלחת, רק רגע...")
    analysis = get_missing_fields_prompt(text)
    
    if not analysis:
        bot.send_message(chat_id, "היה קושי קטן בתקשורת עם השרת, אנא נסה לשלוח שוב.")
        return
        
    missing = analysis.get("missing_fields_keys", [])
    # נגזור את רשימת המפתחות החסרים האמיתיים מתוך ה-extracted_data ששווים ל-null
    extracted_data = analysis.get("extracted_data", {})
    actual_missing = [k for k, v in extracted_data.items() if v is None or v == "null"]
    
    if actual_missing:
        # שמירת המצב במילון
        user_states[chat_id] = {
            "data": extracted_data,
            "missing_list": actual_missing,
            "waiting_for": actual_missing[0]
        }
        
        # מפות שמות ידידותיים לשאלה
        prompts_map = {
            "project_num": "מספר פרויקט", "letter_num": "מספר מכתב",
            "client_name": "שם הלקוח (לכבוד)", "contact_person": "איש קשר (לידי)",
            "email": "כתובת אימייל (במייל)", "structure_name": "שם המבנה / הפרויקט",
            "supervisor_name": "שם המפקח באתר", "contractor_name": "נציגי הביצוע / קבלן",
            "author_initials": "ראשי תיבות של כותב הדוח", "description": "תיאור מצב העבודה הנוכחי"
        }
        
        first_missing = actual_missing[0]
        field_display = prompts_map.get(first_missing, first_missing)
        bot.send_message(chat_id, f"הבנתי את רוב הפרטים, אבל חסרים לי כמה נתונים.\nמהו **{field_display}**?")
    else:
        # הכל קיים כבר בטקסט המקורי!
        bot.send_message(chat_id, "איזה יופי, כל הפרטים קיימים! מפיק את קובץ הוורד...")
        try:
            doc_file = create_report(extracted_data)
            with open(doc_file, 'rb') as f:
                bot.send_document(chat_id, f, caption="הנה דוח הפיקוח העליון המוכן שלך! 📝")
        except Exception as e:
            bot.send_message(chat_id, f"אירעה שגיאה ביצירת הקובץ: {e}")

if __name__ == '__main__':
    print("Bot is running...")
    bot.infinity_polling()